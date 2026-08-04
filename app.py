import os
import json
import sqlite3
from datetime import datetime
from io import BytesIO

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st

APP_NAME = "SITRA-Campamentos"
APP_FULL = "Sistema Inteligente de Transiliencia, Riesgo y Atención para Campamentos Transitorios"
APP_VERSION = "v0.6.1"
DB_PATH = os.path.join("data", "sitra_campamentos.db")

st.set_page_config(page_title=APP_NAME, page_icon="🛡️", layout="wide")

CSS = """
<style>
:root {
  --navy:#0B1F3A;
  --blue:#0B5CAD;
  --sky:#DDF3FF;
  --yellow:#FFD447;
  --red:#D62828;
  --soft:#F7FBFF;
  --gray:#5F6B7A;
  --green:#2A9D8F;
  --orange:#F77F00;
}
.main { background: #ffffff; }
.block-container { padding-top: 1.1rem; padding-bottom: 2rem; }
.hero {
  background: linear-gradient(135deg,#ffffff 0%,#F0FAFF 72%,#E2F5FF 100%);
  border: 1px solid #CBE9F8;
  border-radius: 22px;
  padding: 1.35rem 1.45rem;
  box-shadow: 0 4px 16px rgba(11,31,58,.08);
  position: relative;
  overflow: hidden;
}
.hero:before {
  content:"";
  position:absolute;
  left:0; top:0; height:6px; width:100%;
  background: linear-gradient(90deg, var(--yellow) 0 33%, var(--blue) 33% 66%, var(--red) 66% 100%);
}
.hero h1 { color: var(--navy); margin: .25rem 0 .2rem 0; font-size: 2.2rem; }
.hero h3 { color: var(--blue); margin: .2rem 0 .7rem 0; font-weight: 500; }
.badge {
  display:inline-block; padding:.25rem .55rem; margin:.12rem; border-radius: 999px;
  background:#EAF7FF; color:#0B5CAD; border:1px solid #BFE6FB; font-size:.78rem;
}
.card {
  background:#ffffff; border:1px solid #E6EEF5; border-radius:18px; padding:1rem; 
  box-shadow: 0 2px 10px rgba(11,31,58,.05); margin-bottom:.85rem;
}
.card h4 { color:var(--navy); margin-top:0; }
.notice { background:#EAF4FF; border-left:5px solid var(--blue); padding:.85rem; border-radius:10px; }
.warn { background:#FFF7E1; border-left:5px solid #FFB703; padding:.85rem; border-radius:10px; }
.danger { background:#FFF0F0; border-left:5px solid var(--red); padding:.85rem; border-radius:10px; }
.ok { background:#ECFFF5; border-left:5px solid var(--green); padding:.85rem; border-radius:10px; }
.small { color:var(--gray); font-size:.86rem; }
.metric-card { background:#F8FBFF; border:1px solid #E2EEF9; border-radius:18px; padding:1rem; }
.footer { border-top:1px solid #E8EEF4; padding-top:1rem; color:#6D7785; font-size:.82rem; }
[data-testid="stSidebar"] { background:#F4F8FC; }
.stToolbarActions{ display: none; }
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)

RISK_CATEGORIES = [
    "Vida y seguridad estructural",
    "Salud y apoyo psicosocial",
    "Agua, higiene y saneamiento",
    "Alimentación y suministros",
    "Convivencia y protección",
    "Información y gobernanza",
    "Logística y continuidad",
    "Entorno y riesgos emergentes",
]


def data_path(name):
    return os.path.join("data", name)


def ensure_db():
    os.makedirs("data", exist_ok=True)
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
    CREATE TABLE IF NOT EXISTS diagnosticos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT,
        campamento TEXT,
        municipio TEXT,
        responsable TEXT,
        data_json TEXT,
        riesgos_json TEXT,
        transiliencia_json TEXT,
        plan_json TEXT,
        estrategia_json TEXT
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS beneficios (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT,
        campamento TEXT,
        beneficio TEXT,
        grupo_beneficiario TEXT,
        accion TEXT,
        indicador TEXT,
        linea_base TEXT,
        meta TEXT,
        resultado TEXT,
        estado TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS emprendimientos_agua (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT,
        campamento TEXT,
        nombre_iniciativa TEXT,
        modelo TEXT,
        producto_servicio TEXT,
        talento TEXT,
        aliado TEXT,
        indicador TEXT,
        proxima_accion TEXT,
        etapa TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS capacitaciones_wash (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT,
        campamento TEXT,
        facilitador TEXT,
        participante_grupo TEXT,
        modulo TEXT,
        personas_formadas INTEGER,
        practica_realizada TEXT,
        resultado TEXT,
        proxima_accion TEXT
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS evaluaciones_wash_flash (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT,
        campamento TEXT,
        grupo TEXT,
        modulo TEXT,
        participantes INTEGER,
        evaluacion_entrada REAL,
        evaluacion_salida REAL,
        practica TEXT,
        microcredencial TEXT,
        observaciones TEXT,
        proxima_accion TEXT
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS compromisos_wash (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT,
        campamento TEXT,
        brigada TEXT,
        compromiso TEXT,
        responsable TEXT,
        indicador TEXT,
        plazo TEXT,
        estado TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS analisis_agua (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT,
        campamento TEXT,
        punto TEXT,
        crl REAL,
        turbidez REAL,
        coliformes TEXT,
        apto TEXT,
        analista TEXT,
        certificacion TEXT,
        accion_correctiva TEXT
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS gerencia_centro (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT,
        campamento TEXT,
        gerente TEXT,
        suplente TEXT,
        organizaciones TEXT,
        necesidad_prioritaria TEXT,
        plan_24h TEXT,
        plan_72h TEXT,
        estado TEXT
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS demanda_curso_wash (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT,
        campamento TEXT,
        personas_objetivo INTEGER,
        perfil_objetivo TEXT,
        razon_necesidad TEXT,
        curso_requerido TEXT,
        estado TEXT,
        proxima_accion TEXT
    )
    """)
    con.commit(); con.close()


def load_csv(name):
    path = data_path(name)
    if os.path.exists(path):
        # CSVs are generated with UTF-8 BOM and quoted fields to protect commas in Spanish text.
        return pd.read_csv(path, encoding="utf-8-sig")
    return pd.DataFrame()


def score_level(score):
    if score >= 75: return "Rojo", "Crítico", "#D62828"
    if score >= 55: return "Naranja", "Alto", "#F77F00"
    if score >= 35: return "Amarillo", "Medio", "#F4B400"
    return "Verde", "Controlado", "#2A9D8F"


def clamp(x, a=0, b=100):
    return max(a, min(b, x))


def calculate_risks(d):
    total = max(1, int(d.get("personas_total", 1)))
    cap = max(1, int(d.get("capacidad", 1)))
    occupancy = (total / cap) * 100
    water_liters = float(d.get("agua_litros", 0))
    water_pp = water_liters / total
    bathrooms = max(0, int(d.get("banos", 0)))
    people_per_bath = total / max(1, bathrooms)

    structural = int(d.get("danio_estructural", 0))
    replicas = int(d.get("replicas", 0))
    security = clamp(structural*22 + replicas*12 + max(0, occupancy-100)*0.8)

    critical_med = int(d.get("casos_medicos_criticos", 0))
    sick = int(d.get("personas_enfermas", 0))
    psychosocial = int(d.get("tension_psicosocial", 0))
    health = clamp(critical_med*16 + sick/total*100*1.2 + psychosocial*12)

    waste = int(d.get("residuos", 0))
    crl = float(d.get("crl_consumo", 0))
    turbidez = float(d.get("turbidez_ntu", 0))
    lecturas_dia = int(d.get("lecturas_dia", 0))
    coliformes = int(d.get("coliformes_detectados", 0))
    certified = int(d.get("certificados_agua", 0))
    need_course = int(d.get("necesidad_curso_agua", 0))
    water_quantity_risk = clamp((15-water_pp)*5 if water_pp < 15 else 0)
    sanitation_risk = clamp((people_per_bath-20)*1.2 if people_per_bath > 20 else 0) + waste*15
    quality_risk = 0
    if crl < 0.2: quality_risk += 28
    elif crl < 0.5: quality_risk += 16
    elif crl > 1.5: quality_risk += 16
    if turbidez > 5: quality_risk += 22
    if coliformes: quality_risk += 35
    if lecturas_dia < 3: quality_risk += 12
    if certified <= 0: quality_risk += 12
    if need_course >= 3: quality_risk += 10
    wash = clamp(water_quantity_risk + sanitation_risk + quality_risk)

    food_days = float(d.get("dias_alimentos", 0))
    supply = int(d.get("cadena_suministro", 0))
    food = clamp((3-food_days)*20 if food_days < 3 else 0) + supply*18
    food = clamp(food)

    conflicts = int(d.get("conflictos", 0))
    vuln = int(d.get("vulnerables", 0))
    protection = clamp(conflicts*18 + (vuln/total)*100*0.6)

    rumors = int(d.get("rumores", 0))
    official_info = int(d.get("info_oficial", 0))
    governance = clamp(rumors*20 + max(0, 4-official_info)*10)

    energy = int(d.get("energia", 0))
    access = int(d.get("acceso", 0))
    inventory = int(d.get("inventario", 0))
    logistics = clamp(energy*12 + access*16 + inventory*14)

    rain = int(d.get("lluvia", 0))
    emerging = int(d.get("senales_tempranas", 0))
    external = int(d.get("apoyo_externo", 0))
    environment = clamp(rain*14 + emerging*18 + max(0, 4-external)*8)

    values = {
        "Vida y seguridad estructural": security,
        "Salud y apoyo psicosocial": health,
        "Agua, higiene y saneamiento": wash,
        "Alimentación y suministros": food,
        "Convivencia y protección": protection,
        "Información y gobernanza": governance,
        "Logística y continuidad": logistics,
        "Entorno y riesgos emergentes": environment,
    }
    rows = []
    for cat, val in values.items():
        color_name, level, color = score_level(val)
        rows.append({"Categoría": cat, "Puntaje": round(val,1), "Semáforo": color_name, "Nivel": level, "Color": color})
    return pd.DataFrame(rows)


def calculate_strategy_status(d):
    # Madurez breve por los 12 elementos estratégicos a partir de campos del diagnóstico.
    status = {
        "EE-01": int(d.get("registros",0))*18 + int(d.get("comites",0))*8,
        "EE-02": int(d.get("apoyo_entorno",0))*18 + int(d.get("actores_entorno",0))*8,
        "EE-03": int(d.get("apoyo_extorno",0))*18,
        "EE-04": int(d.get("destrezas_extorno",0))*18,
        "EE-05": int(d.get("voluntarios",0))/max(1,int(d.get("personas_total",1)))*200 + int(d.get("microservicios",0))*14,
        "EE-06": max(0,4-int(d.get("inventario",0)))*18,
        "EE-07": int(d.get("info_oficial",0))*14 + int(d.get("apoyo_entorno",0))*10,
        "EE-08": int(d.get("bloqueos_mitigados",0))*20,
        "EE-09": int(d.get("erp",0))*25,
        "EE-10": int(d.get("registros",0))*20 + int(d.get("talento_mapeado",0))*15,
        "EE-11": int(d.get("normas_donantes",0))*22,
        "EE-12": int(d.get("buenas_practicas",0))*22,
    }
    rows=[]
    for k,v in status.items():
        color, level, _ = score_level(clamp(v))
        rows.append({"Código":k,"Madurez":round(clamp(v),1),"Nivel":level,"Semáforo":color})
    return pd.DataFrame(rows)


def generate_plan(risks_df, d, strategy_df):
    objectives = load_csv("objetivos_sistemicos.csv")
    elements = load_csv("elementos_estrategicos.csv")
    actions = []
    templates = {
        "Vida y seguridad estructural": ("Delimitar zonas inseguras, verificar grietas visibles, ubicar punto seguro y escalar a Protección Civil/ingeniería si hay daño estructural o réplicas significativas.", "OE-1.2", "EE-09"),
        "Salud y apoyo psicosocial": ("Priorizar casos críticos, activar registro sanitario, separar zona de primeros auxilios y solicitar apoyo médico/psicosocial competente.", "OE-1.1", "EE-10"),
        "Agua, higiene y saneamiento": ("Implementar racionamiento transparente, puntos de lavado de manos, limpieza por turnos, disposición de residuos y solicitud urgente de agua/baños portátiles si aplica.", "OE-1.2", "EE-06"),
        "Alimentación y suministros": ("Calcular autonomía real, ordenar inventario, definir turnos de distribución visibles y escalar necesidades de alimentos para 72 horas.", "OE-1.2", "EE-06"),
        "Convivencia y protección": ("Crear comité de convivencia, mediación, canales de quejas, protección de niños/adultos mayores/personas con discapacidad y vigilancia comunitaria no violenta.", "OE-1.1", "EE-05"),
        "Información y gobernanza": ("Designar vocero único, cartelera oficial, boletín diario, registro de rumores y mecanismo de verificación antes de difundir información.", "OE-2.2", "EE-07"),
        "Logística y continuidad": ("Identificar recursos críticos, responsables por área, rutas de acceso, energía alternativa, comunicaciones y prioridades de continuidad del campamento transitorio.", "OE-2.1", "EE-09"),
        "Entorno y riesgos emergentes": ("Activar radar de señales tempranas: clima, réplicas, seguridad externa, abastecimiento, tensiones comunitarias y cambios de contexto.", "OE-3.2", "EE-08"),
    }
    for _, r in risks_df.sort_values("Puntaje", ascending=False).iterrows():
        if r["Puntaje"] >= 55:
            horizon = "0-24 h" if r["Puntaje"] >= 75 else "24-72 h"
            text, obj, elem = templates[r["Categoría"]]
            obj_text = objectives.loc[objectives["Codigo"]==obj,"Objetivo"].iloc[0] if not objectives.empty and obj in objectives["Codigo"].values else obj
            elem_text = elements.loc[elements["Codigo"]==elem,"Elemento"].iloc[0] if not elements.empty and elem in elements["Codigo"].values else elem
            actions.append({
                "Horizonte": horizon,
                "Prioridad": r["Nivel"],
                "Categoría": r["Categoría"],
                "Acción recomendada": text,
                "Objetivo relacionado": obj,
                "Elemento estratégico": elem,
                "Razón sistémica": f"Contribuye a {obj}: {obj_text[:120]}... | Activa {elem}: {elem_text[:100]}...",
                "Responsable sugerido": "Coordinador del campamento + líder de brigada",
                "Evidencia": "Foto / registro / acta / lista / reporte"
            })

    # Acción productiva especial para agua potable y WASH
    wash_rows = risks_df[risks_df["Categoría"] == "Agua, higiene y saneamiento"]
    if not wash_rows.empty and float(wash_rows.iloc[0]["Puntaje"]) >= 35:
        actions.append({
            "Horizonte": "24-72 h",
            "Prioridad": "Productiva-WASH",
            "Categoría": "Agua potable productiva",
            "Acción recomendada": "Activar la Brigada Agua Segura: medir cloro residual libre, turbidez observable, puntos de consumo y necesidades de cloración; convertir el control WASH en microservicio comunitario trazable.",
            "Objetivo relacionado": "OE-1.2",
            "Elemento estratégico": "EE-10",
            "Razón sistémica": "Convierte el mapa de talento en capacidad operativa para asegurar agua potable, higiene y salud pública mientras se reconstruye La Guaira.",
            "Responsable sugerido": "Coordinador del campamento + Brigada Agua Segura + enlace sanitario",
            "Evidencia": "Bitácora WASH, lecturas CRL, fotos de puntos de agua, lista de brigadistas"
        })
        actions.append({
            "Horizonte": "7-30 días",
            "Prioridad": "Emprendimiento",
            "Categoría": "Productivización y servitización del agua",
            "Acción recomendada": "Diseñar un emprendimiento comunitario de agua potable: producto mínimo viable, servicio de monitoreo, capacitación, aliados, indicadores y protocolo de sostenibilidad/sustentabilidad/soportabilidad.",
            "Objetivo relacionado": "OG-General",
            "Elemento estratégico": "EE-06 / EE-09 / EE-11",
            "Razón sistémica": "Pasa de recibir ayuda a crear capacidades y servicios comunitarios que sostienen el agua segura y promueven ciudadanía inteligente.",
            "Responsable sugerido": "Mesa de Agua Potable Productiva",
            "Evidencia": "Lienzo de modelo, inventario, acuerdos, indicadores y plan piloto"
        })

    # Recomendaciones específicas del tutor: el análisis de calidad de agua debe nacer en el diagnóstico inicial,
    # generar necesidad de Escuela WASH y habilitar la evaluación solo con personal entrenado/certificado.
    crl = float(d.get("crl_consumo", 0))
    turbidez = float(d.get("turbidez_ntu", 0))
    lecturas_dia = int(d.get("lecturas_dia", 0))
    coliformes = int(d.get("coliformes_detectados", 0))
    certificados = int(d.get("certificados_agua", 0))
    necesidad_curso = int(d.get("necesidad_curso_agua", 0))
    if crl < 0.2 or turbidez > 5 or coliformes or lecturas_dia < 3:
        actions.append({
            "Horizonte": "0-24 h",
            "Prioridad": "Agua-Crítica",
            "Categoría": "Calidad del agua en diagnóstico inicial",
            "Acción recomendada": "No dejar la calidad del agua para después: registrar CRL, turbidez, coliformes/ausencia de prueba y número de lecturas diarias; activar acción correctiva y escalar a autoridad sanitaria si no se cumplen parámetros mínimos.",
            "Objetivo relacionado": "OE-1.2",
            "Elemento estratégico": "EE-06 / EE-10",
            "Razón sistémica": "El agua es eje de vida, salud, higiene y convivencia; por eso el diagnóstico inicial debe impactar directamente la priorización y el plan de acción.",
            "Responsable sugerido": "Gerente del centro + Monitor de Agua Segura + enlace sanitario",
            "Evidencia": "Bitácora de análisis de agua, foto del punto, lectura CRL/turbidez, acción correctiva"
        })
    if certificados <= 0 or necesidad_curso >= 2:
        actions.append({
            "Horizonte": "0-72 h",
            "Prioridad": "Formación urgente",
            "Categoría": "Escuela WASH - análisis rápido de agua",
            "Acción recomendada": "Generar la necesidad del curso: seleccionar personas con competencias base, entrenarlas en análisis rápido de agua, bitácora y criterios de no liberación; hasta certificar, los resultados deben ser revisados por un responsable técnico o sanitario.",
            "Objetivo relacionado": "OE-1.1 / OE-1.2",
            "Elemento estratégico": "EE-10 / EE-11",
            "Razón sistémica": "Primero se entrena a la gente para analizar el agua; luego se usa esa capacidad para productivizar servicios WASH dentro y fuera del campamento.",
            "Responsable sugerido": "Gerente del centro + Escuela WASH + equipo de biología/ecología/salud",
            "Evidencia": "Lista de inscritos, evaluación práctica, microcredencial, designación de monitor"
        })
    # Agregar acciones para elementos estratégicos débiles
    weak = strategy_df[strategy_df["Madurez"] < 35].head(3)
    for _, w in weak.iterrows():
        elem = w["Código"]
        elem_text = elements.loc[elements["Codigo"]==elem,"Elemento"].iloc[0] if not elements.empty and elem in elements["Codigo"].values else elem
        actions.append({
            "Horizonte": "7-30 días",
            "Prioridad": "Estratégica",
            "Categoría": "12 elementos estratégicos",
            "Acción recomendada": f"Elevar madurez de {elem}: {elem_text}.",
            "Objetivo relacionado": "OG-General",
            "Elemento estratégico": elem,
            "Razón sistémica": "Alinea el campamento con los 12 elementos estratégicos clave y los objetivos del análisis sistémico.",
            "Responsable sugerido": "Coordinación estratégica / enlace institucional",
            "Evidencia": "Bitácora, matriz de actores, inventario o protocolo"
        })
    if not actions:
        actions.append({
            "Horizonte": "7 días",
            "Prioridad": "Preventiva",
            "Categoría": "Mejora continua",
            "Acción recomendada": "Mantener monitoreo diario, simulacro de evacuación, registro de recursos, actualización de historia viva y comunicación con comunidad.",
            "Objetivo relacionado": "OG-General",
            "Elemento estratégico": "EE-10",
            "Razón sistémica": "Consolida habitabilidad digna, sostenibilidad operativa y capacidad de respuesta.",
            "Responsable sugerido": "Coordinador del campamento",
            "Evidencia": "Bitácora diaria"
        })
    return pd.DataFrame(actions)


def calculate_transilience(d):
    total = max(1, int(d.get("personas_total", 1)))
    volunteers = int(d.get("voluntarios", 0))
    committees = int(d.get("comites", 0))
    trainings = int(d.get("capacitaciones", 0))
    microservices = int(d.get("microservicios", 0))
    drills = int(d.get("simulacros", 0))
    records = int(d.get("registros", 0))
    conflicts = int(d.get("conflictos", 0))
    rumors = int(d.get("rumores", 0))
    info = int(d.get("info_oficial", 0))
    buenas = int(d.get("buenas_practicas", 0))

    autonomy = clamp(volunteers/total*100*3 + committees*10)
    cooperation = clamp(committees*18 + max(0, 5-conflicts)*10)
    trust = clamp(max(0, 5-rumors)*12 + info*10)
    learning = clamp(trainings*18 + drills*10 + buenas*8)
    productivity = clamp(microservices*20 + volunteers/total*100*2)
    prevention = clamp(drills*25 + records*8)
    data = clamp(records*18 + info*10)
    smart_citizenship = clamp((autonomy+cooperation+learning+data)/4)

    dims = {
        "Autonomía": autonomy,
        "Cooperación": cooperation,
        "Confianza": trust,
        "Aprendizaje": learning,
        "Productividad comunitaria": productivity,
        "Cultura preventiva": prevention,
        "Trazabilidad y datos": data,
        "Ciudadanía inteligente": smart_citizenship,
    }
    index = round(sum(dims.values())/len(dims),1)
    if index < 25: stage = "R0 Emergencia"
    elif index < 45: stage = "R1 Estabilización"
    elif index < 65: stage = "R2 Resiliencia"
    elif index < 82: stage = "T1 Transiliencia"
    else: stage = "T2 Ciudadanía inteligente"
    return dims, index, stage


def save_diagnostic(d, risks_df, trans, plan_df, strategy_df):
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
    INSERT INTO diagnosticos(fecha, campamento, municipio, responsable, data_json, riesgos_json, transiliencia_json, plan_json, estrategia_json)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        datetime.now().isoformat(timespec="seconds"),
        d.get("campamento", "Sin nombre"),
        d.get("municipio", ""),
        d.get("responsable", ""),
        json.dumps(d, ensure_ascii=False),
        risks_df.to_json(orient="records", force_ascii=False),
        json.dumps(trans, ensure_ascii=False),
        plan_df.to_json(orient="records", force_ascii=False),
        strategy_df.to_json(orient="records", force_ascii=False)
    ))
    con.commit(); con.close()


def last_diagnostic():
    con = sqlite3.connect(DB_PATH)
    df = pd.read_sql_query("SELECT * FROM diagnosticos ORDER BY id DESC LIMIT 1", con)
    con.close()
    if df.empty: return None
    row = df.iloc[0].to_dict()
    row["data"] = json.loads(row["data_json"])
    row["risks"] = pd.DataFrame(json.loads(row["riesgos_json"]))
    row["trans"] = json.loads(row["transiliencia_json"])
    row["plan"] = pd.DataFrame(json.loads(row["plan_json"]))
    row["strategy"] = pd.DataFrame(json.loads(row["estrategia_json"])) if row.get("estrategia_json") else pd.DataFrame()
    return row


def all_diagnostics():
    con = sqlite3.connect(DB_PATH)
    df = pd.read_sql_query("SELECT * FROM diagnosticos ORDER BY id DESC", con)
    con.close()
    return df


def hero():
    st.markdown(f"""
    <div class="hero">
      <h1>{APP_NAME}</h1>
      <h3>{APP_FULL}</h3>
      <p>Sistema experto local para apoyar a gerentes de campamentos transitorios en diagnóstico, priorización de riesgos, gestión de beneficios, FODA sistémica, objetivos estratégicos, plan de acción y ruta de transiliencia comunitaria.</p>
      <span class="badge">Campamentos Transitorios</span><span class="badge">ISO 31000</span><span class="badge">IEC 31010</span><span class="badge">ISO 22316</span><span class="badge">FODA Sistémica</span><span class="badge">12 Elementos Estratégicos</span><span class="badge">Smart City</span><span class="badge">Océano Azul</span><span class="badge">WASH</span><span class="badge">Agua potable productiva</span>
    </div>
    """, unsafe_allow_html=True)


def page_inicio():
    hero()
    st.write("")
    c1,c2,c3,c4 = st.columns(4)
    refs = load_csv("referencias.csv")
    c1.metric("Categorías de riesgo", len(RISK_CATEGORIES))
    c2.metric("Elementos estratégicos", 12)
    c3.metric("Objetivos sistémicos", 7)
    c4.metric("Fuentes visibles", len(refs) if not refs.empty else 0)
    st.subheader("Qué evalúa y organiza")
    st.markdown("""
    - **Riesgos críticos** del campamento transitorio: vida, salud, agua, alimentación, convivencia, información, logística y entorno.
    - **FODA sistémica** en tres escalas: Intorno, Entorno y Extorno.
    - **12 elementos estratégicos** para sostener, articular, inventariar, mapear talento, agilizar donaciones y aprender de experiencias sísmicas acertadas.
    - **Objetivos del análisis sistémico**, conectando cada decisión con habitabilidad digna, sostenibilidad operativa y capacidad de respuesta.
    - **Transiliencia ciudadana**, para pasar de asistencia pasiva a comunidad organizada, productiva, servicial y preparada para reconstrucción inteligente.
    - **Agua potable productiva**, para orientar brigadas y emprendimientos comunitarios basados en WASH, productivización y servitización del agua segura en La Guaira.
    - **Escuela WASH Flash**, para entrenar rapidamente a dignificados, brigadas y lideres comunitarios en agua segura, saneamiento, higiene, bioseguridad y emprendimiento comunitario.
    """)
    st.subheader("Manual de usuario")
    st.markdown("Descargue el manual para explicar el alcance de cada modulo y usar SITRA como acelerador de decisiones, estrategias y acciones verificables.")
    manual_docx = os.path.join("docs", "Manual_Usuario_SITRA_Campamentos_v0_6_1.docx")
    manual_pdf = os.path.join("docs", "Manual_Usuario_SITRA_Campamentos_v0_6_1.pdf")
    cdoc, cpdf = st.columns(2)
    if os.path.exists(manual_docx):
        with open(manual_docx, "rb") as f:
            cdoc.download_button("Descargar manual Word", f.read(), "Manual_Usuario_SITRA_Campamentos_v0_6_1.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    if os.path.exists(manual_pdf):
        with open(manual_pdf, "rb") as f:
            cpdf.download_button("Descargar manual PDF", f.read(), "Manual_Usuario_SITRA_Campamentos_v0_6_1.pdf", "application/pdf")
    st.markdown("""
    <div class="notice"><b>Uso responsable:</b> esta herramienta apoya decisiones. No sustituye a Protección Civil, autoridades sanitarias, bomberos, equipos médicos, equipos psicosociales ni organismos oficiales.</div>
    """, unsafe_allow_html=True)
    st.write("")
    st.subheader("Acceso desde Android")
    st.markdown("""
    1. Ejecuta el sistema en la PC con <b>INICIAR_SITRA_CAMPAMENTOS.bat</b>.
    2. PC y Android deben estar en la misma red WiFi.
    3. En la PC, ejecuta <b>VER_IP_DE_ESTA_PC.bat</b> para ver la IP.
    4. En Android abre: <code>http://IP_DE_LA_PC:8501</code>.
    """, unsafe_allow_html=True)


def page_referencias():
    st.title("Base metodológica y transparencia")
    st.markdown("Estas fuentes son visibles para explicar por qué la herramienta recomienda, prioriza y organiza decisiones.")
    df = load_csv("referencias.csv")
    if df.empty:
        st.warning("No se encontró referencias.csv")
        return
    for _, r in df.iterrows():
        link = r['Enlace']
        link_html = f'<a href="{link}" target="_blank">Abrir fuente verificable</a>' if str(link).startswith("http") else f'<span class="small">{link}</span>'
        st.markdown(f"""
        <div class="card">
          <h4>{r['Fuente']}</h4>
          <p><b>Uso:</b> {r['Uso']}</p>
          <p><b>Aporte en SITRA-Campamentos:</b> {r['Aporte_SECP']}</p>
          <p><b>Tipo:</b> {r['Tipo']} &nbsp; | &nbsp; <b>Autoridad:</b> {r['Autoridad']}</p>
          <p>{link_html}</p>
        </div>
        """, unsafe_allow_html=True)
    st.markdown("""
    <div class="warn"><b>Nota metodológica:</b> Smart City se usa aquí como marco sociocultural y sistémico: agentes, redes, conductas colectivas, espacio e información. Océano Azul se usa como lógica de innovación de valor humanitaria: no competir por ayuda dispersa, sino crear una red diferenciada de beneficios, coordinación y capacidades.</div>
    """, unsafe_allow_html=True)


def page_diagnostico():
    st.title("Diagnóstico rápido del campamento transitorio")
    st.caption("Complete el formulario con la mejor información disponible. En emergencia, dato imperfecto pero trazable vale más que silencio perfecto.")
    with st.form("diag_form"):
        st.subheader("Identificación")
        c1,c2,c3 = st.columns(3)
        campamento = c1.text_input("Nombre del campamento transitorio", "Campamento Demo Libertador")
        municipio = c2.text_input("Municipio / parroquia", "")
        responsable = c3.text_input("Responsable", "")
        st.subheader("Población y capacidad")
        c1,c2,c3,c4 = st.columns(4)
        personas_total = c1.number_input("Personas totales", min_value=1, value=180, step=1)
        capacidad = c2.number_input("Capacidad estimada", min_value=1, value=150, step=1)
        vulnerables = c3.number_input("Personas vulnerables", min_value=0, value=45, step=1)
        voluntarios = c4.number_input("Voluntarios activos", min_value=0, value=18, step=1)
        st.subheader("Recursos críticos")
        c1,c2,c3,c4 = st.columns(4)
        agua_litros = c1.number_input("Agua disponible total (litros/día)", min_value=0, value=1800, step=50)
        banos = c2.number_input("Baños operativos", min_value=0, value=6, step=1)
        dias_alimentos = c3.number_input("Días de alimentos disponibles", min_value=0.0, value=1.5, step=0.5)
        cadena_suministro = c4.slider("Tensión cadena de suministro", 0, 4, 2, help="0 estable, 4 crítica")
        st.subheader("Calidad del agua y necesidad de análisis WASH")
        st.caption("Requerimiento del tutor: la evaluación de calidad del agua debe estar en el diagnóstico inicial para impactar riesgo, plan de acción y necesidad de Escuela WASH.")
        c1,c2,c3,c4 = st.columns(4)
        crl_consumo = c1.number_input("CRL en punto de consumo (mg/L)", min_value=0.0, max_value=5.0, value=0.2, step=0.1, help="Cloro residual libre observado. Meta operativa de emergencia: 0.2-0.5 mg/L; en bitácora ampliada puede manejarse 0.5-1.0 según contexto.")
        turbidez_ntu = c2.number_input("Turbidez estimada/medida (NTU)", min_value=0.0, max_value=100.0, value=6.0, step=0.5, help="Meta técnica de referencia: < 5 NTU cuando se dispone de medición.")
        lecturas_dia = c3.number_input("Lecturas de agua por día", min_value=0, max_value=12, value=1, step=1, help="SITRA recomienda mínimo 3 puntos/lecturas diarias en emergencia.")
        coliformes_detectados = c4.selectbox("Coliformes fecales detectados o sospechados", [0,1], format_func=lambda x: "No / sin evidencia" if x==0 else "Sí / sospecha / sin prueba confiable")
        c1,c2,c3,c4 = st.columns(4)
        certificados_agua = c1.number_input("Personas certificadas para análisis de agua", min_value=0, value=0, step=1)
        necesidad_curso_agua = c2.slider("Necesidad de curso rápido de análisis de agua", 0, 4, 4, help="0 sin necesidad; 4 urgente.")
        curso_agua_estado = c3.selectbox("Estado del curso de análisis de agua", ["No iniciado", "Programado", "En curso", "Certificación pendiente", "Personas certificadas"])
        analisis_agua_habilitado = c4.selectbox("¿Usar evaluación formal de agua en decisiones?", ["Solo orientación", "Habilitado con responsable técnico", "Habilitado con personal certificado"])

        st.subheader("Riesgos y señales")
        c1,c2,c3,c4 = st.columns(4)
        danio_estructural = c1.slider("Daño estructural visible", 0, 4, 2)
        replicas = c2.slider("Riesgo/percepción de réplicas", 0, 4, 2)
        personas_enfermas = c3.number_input("Personas enfermas", min_value=0, value=12, step=1)
        casos_medicos_criticos = c4.number_input("Casos médicos críticos", min_value=0, value=2, step=1)
        c1,c2,c3,c4 = st.columns(4)
        tension_psicosocial = c1.slider("Tensión psicosocial colectiva", 0, 4, 3)
        conflictos = c2.slider("Conflictos/convivencia", 0, 4, 2)
        rumores = c3.slider("Rumores/desinformación", 0, 4, 2)
        info_oficial = c4.slider("Comunicación oficial efectiva", 0, 4, 2, help="0 ausente, 4 muy efectiva")
        c1,c2,c3,c4 = st.columns(4)
        residuos = c1.slider("Problemas de residuos/saneamiento", 0, 4, 2)
        energia = c2.slider("Falla de energía/iluminación", 0, 4, 2)
        acceso = c3.slider("Dificultad de acceso/transporte", 0, 4, 2)
        inventario = c4.slider("Debilidad de inventario/logística", 0, 4, 2)
        c1,c2,c3 = st.columns(3)
        lluvia = c1.slider("Lluvia/deslizamiento/inundación", 0, 4, 1)
        senales_tempranas = c2.slider("Señales tempranas preocupantes", 0, 4, 2)
        apoyo_externo = c3.slider("Apoyo externo disponible", 0, 4, 2, help="0 ausente, 4 alto")
        st.subheader("Madurez estratégica y transiliencia")
        c1,c2,c3,c4 = st.columns(4)
        comites = c1.number_input("Comités/brigadas activas", min_value=0, value=3, step=1)
        capacitaciones = c2.number_input("Capacitaciones breves realizadas", min_value=0, value=1, step=1)
        microservicios = c3.number_input("Microservicios internos activos", min_value=0, value=2, step=1)
        simulacros = c4.number_input("Simulacros/rutinas preventivas", min_value=0, value=0, step=1)
        c1,c2,c3,c4 = st.columns(4)
        registros = c1.slider("Historia viva / registros", 0, 4, 2)
        talento_mapeado = c2.slider("Mapa de talento y conocimiento", 0, 4, 1)
        apoyo_entorno = c3.slider("Articulación con entorno", 0, 4, 1)
        apoyo_extorno = c4.slider("Articulación con extorno", 0, 4, 0)
        c1,c2,c3,c4 = st.columns(4)
        actores_entorno = c1.slider("Actores del entorno mapeados", 0, 4, 1)
        destrezas_extorno = c2.slider("Destrezas del extorno mapeadas", 0, 4, 0)
        normas_donantes = c3.slider("Normas de donantes conocidas", 0, 4, 0)
        buenas_practicas = c4.slider("Buenas prácticas sísmicas incorporadas", 0, 4, 0)
        c1,c2 = st.columns(2)
        bloqueos_mitigados = c1.slider("Rutas ante bloqueos/condiciones externas", 0, 4, 0)
        erp = c2.slider("Soporte ERP / sistema de eventos", 0, 4, 1)
        submitted = st.form_submit_button("Calcular y guardar diagnóstico")
    if submitted:
        d = locals().copy()
        allowed = ["campamento","municipio","responsable","personas_total","capacidad","vulnerables","voluntarios","agua_litros","banos","dias_alimentos","cadena_suministro","danio_estructural","replicas","personas_enfermas","casos_medicos_criticos","tension_psicosocial","conflictos","rumores","info_oficial","residuos","energia","acceso","inventario","lluvia","senales_tempranas","apoyo_externo","comites","capacitaciones","microservicios","simulacros","registros","talento_mapeado","apoyo_entorno","apoyo_extorno","actores_entorno","destrezas_extorno","normas_donantes","buenas_practicas","bloqueos_mitigados","erp","crl_consumo","turbidez_ntu","lecturas_dia","coliformes_detectados","certificados_agua","necesidad_curso_agua","curso_agua_estado","analisis_agua_habilitado"]
        d = {k:d[k] for k in allowed}
        risks = calculate_risks(d)
        strategy = calculate_strategy_status(d)
        dims, idx, stage = calculate_transilience(d)
        trans = {"dimensiones": dims, "indice": idx, "etapa": stage}
        plan = generate_plan(risks, d, strategy)
        save_diagnostic(d, risks, trans, plan, strategy)
        st.success("Diagnóstico guardado correctamente.")
        show_results(risks, trans, plan, strategy)


def show_results(risks, trans, plan, strategy):
    st.subheader("Resultado del diagnóstico")
    c1,c2,c3,c4 = st.columns(4)
    avg = round(risks["Puntaje"].mean(),1)
    _, level, _ = score_level(avg)
    c1.metric("Riesgo promedio", avg, level)
    c2.metric("Índice de transiliencia", trans["indice"], trans["etapa"])
    c3.metric("Madurez estratégica", round(strategy["Madurez"].mean(),1))
    c4.metric("Acciones recomendadas", len(plan))
    fig = px.bar(risks, x="Categoría", y="Puntaje", color="Nivel", text="Puntaje", title="Mapa de riesgos por categoría")
    fig.update_layout(xaxis_tickangle=-35, height=430, showlegend=True)
    st.plotly_chart(fig, use_container_width=True)
    st.dataframe(plan, use_container_width=True, hide_index=True)


def page_foda():
    st.title("FODA Sistémica · Intorno, Entorno y Extorno")
    df = load_csv("foda_sistemica.csv")
    st.markdown("La FODA Sistémica permite leer el campamento transitorio como sistema abierto: lo interno, lo comunitario-institucional cercano y las fuerzas externas de cooperación o bloqueo.")
    eje = st.selectbox("Eje", ["Todos"] + sorted(df["Eje"].unique().tolist()))
    tipo = st.selectbox("Tipo", ["Todos"] + ["Fortaleza","Debilidad","Oportunidad","Amenaza"])
    view = df.copy()
    if eje != "Todos": view = view[view["Eje"]==eje]
    if tipo != "Todos": view = view[view["Tipo"]==tipo]
    st.dataframe(view, use_container_width=True, hide_index=True)
    pivot = df.pivot_table(index="Eje", columns="Tipo", values="Aspecto", aggfunc="count", fill_value=0).reset_index()
    fig = px.bar(pivot, x="Eje", y=[c for c in pivot.columns if c != "Eje"], barmode="group", title="Conteo FODA por eje sistémico")
    st.plotly_chart(fig, use_container_width=True)


def page_elementos():
    st.title("12 Elementos Estratégicos Clave")
    df = load_csv("elementos_estrategicos.csv")
    st.markdown("Estos 12 elementos funcionan como brújula estratégica: cada decisión del gerente debe contribuir a sostener el campamento, activar talento, conectar actores, gestionar insumos, agilizar donaciones y aprender de experiencias sísmicas acertadas.")
    st.dataframe(df, use_container_width=True, hide_index=True)
    counts = df.groupby("Eje").size().reset_index(name="Cantidad")
    fig = px.pie(counts, names="Eje", values="Cantidad", title="Distribución de elementos por eje")
    st.plotly_chart(fig, use_container_width=True)
    last = last_diagnostic()
    if last and not last["strategy"].empty:
        st.subheader("Madurez estratégica del último diagnóstico")
        st.dataframe(last["strategy"], use_container_width=True, hide_index=True)
        fig2 = px.bar(last["strategy"], x="Código", y="Madurez", color="Nivel", title="Madurez por elemento estratégico")
        st.plotly_chart(fig2, use_container_width=True)


def page_objetivos():
    st.title("Objetivos del análisis sistémico")
    df = load_csv("objetivos_sistemicos.csv")
    general = df[df["Tipo"]=="General"]
    if not general.empty:
        st.markdown(f"<div class='notice'><b>Objetivo General:</b> {general.iloc[0]['Objetivo']}</div>", unsafe_allow_html=True)
    st.subheader("Objetivos específicos por eje")
    st.dataframe(df[df["Tipo"]!="General"], use_container_width=True, hide_index=True)
    counts = df[df["Tipo"]!="General"].groupby("Eje").size().reset_index(name="Objetivos")
    fig = px.bar(counts, x="Eje", y="Objetivos", title="Objetivos específicos por eje")
    st.plotly_chart(fig, use_container_width=True)


def page_mapa_riesgos():
    st.title("Mapa de riesgos")
    last = last_diagnostic()
    if not last:
        st.info("Primero registre un diagnóstico rápido.")
        return
    risks = last["risks"]
    st.markdown(f"<div class='card'><b>Campamento:</b> {last['campamento']} &nbsp; | &nbsp; <b>Fecha:</b> {last['fecha']}</div>", unsafe_allow_html=True)
    fig = px.bar(risks.sort_values("Puntaje", ascending=True), y="Categoría", x="Puntaje", color="Nivel", orientation="h", title="Prioridad de riesgos")
    st.plotly_chart(fig, use_container_width=True)
    st.dataframe(risks[["Categoría","Puntaje","Nivel","Semáforo"]].sort_values("Puntaje", ascending=False), use_container_width=True, hide_index=True)


def page_plan_accion():
    st.title("Plan de acción alineado a objetivos")
    last = last_diagnostic()
    if not last:
        st.info("Primero registre un diagnóstico rápido.")
        return
    plan = last["plan"]
    st.dataframe(plan, use_container_width=True, hide_index=True)
    st.download_button("Descargar plan CSV", plan.to_csv(index=False).encode("utf-8-sig"), "plan_accion_sitra_campamentos.csv", "text/csv")


def page_beneficios_oceano():
    st.title("Gestión de beneficios y enfoque Océano Azul")
    st.markdown("""
    La gestión de beneficios conecta cada acción con un cambio verificable. El enfoque Océano Azul se usa aquí como metáfora estratégica responsable: **crear valor humanitario nuevo**, no competir por la misma ayuda escasa. El objetivo es convertir asistencia dispersa en una red de servicios, capacidades y beneficios medibles.
    """)
    st.markdown("<div class='notice'><b>Clave operativa:</b> aumentar valor para dignificados y gestores, reducir fricción burocrática, eliminar duplicidades y crear nuevos servicios comunitarios dentro del campamento.</div>", unsafe_allow_html=True)
    benefits_catalog = load_csv("beneficios_base.csv")
    if not benefits_catalog.empty:
        st.subheader("Beneficios sugeridos")
        st.dataframe(benefits_catalog, use_container_width=True, hide_index=True)
    st.subheader("Lienzo ERIC humanitario")
    eric = pd.DataFrame({
        "Acción estratégica":["Eliminar","Reducir","Incrementar","Crear"],
        "Pregunta guía":["¿Qué trámites, duplicidades o prácticas no aportan valor?","¿Qué fricciones, rumores, esperas o pérdidas deben bajar?","¿Qué capacidades, transparencia y servicios deben subir?","¿Qué servicios/brigadas/alianzas nuevas debemos crear?"],
        "Ejemplo en campamento":["Listas duplicadas y solicitudes sin trazabilidad","Colas desordenadas y dependencia externa","Historia viva, inventario y vocería oficial","Mapa de talento, mesa de cooperación y microservicios comunitarios"]
    })
    st.dataframe(eric, use_container_width=True, hide_index=True)
    with st.form("benefit_form"):
        st.subheader("Registrar beneficio")
        c1,c2 = st.columns(2)
        campamento = c1.text_input("Campamento transitorio", "Campamento Demo Libertador")
        beneficio = c2.selectbox("Beneficio", benefits_catalog["Beneficio"].tolist() if not benefits_catalog.empty else ["Seguridad", "Salud", "Confianza"])
        grupo = st.text_input("Grupo beneficiario", "Familias del campamento")
        accion = st.text_area("Acción que produce el beneficio", "Crear turnos visibles y comité comunitario de distribución.")
        indicador = st.text_input("Indicador", "Reducción de quejas/conflictos por distribución")
        c1,c2,c3 = st.columns(3)
        linea = c1.text_input("Línea base", "Quejas diarias: 10")
        meta = c2.text_input("Meta", "Quejas diarias < 5")
        resultado = c3.text_input("Resultado observado", "Pendiente")
        estado = st.selectbox("Estado", ["Planificado", "En ejecución", "Logrado", "Requiere apoyo"])
        ok = st.form_submit_button("Guardar beneficio")
    if ok:
        con = sqlite3.connect(DB_PATH)
        cur = con.cursor()
        cur.execute("""INSERT INTO beneficios(fecha,campamento,beneficio,grupo_beneficiario,accion,indicador,linea_base,meta,resultado,estado)
        VALUES(?,?,?,?,?,?,?,?,?,?)""", (datetime.now().isoformat(timespec="seconds"),campamento,beneficio,grupo,accion,indicador,linea,meta,resultado,estado))
        con.commit(); con.close()
        st.success("Beneficio guardado.")
    con = sqlite3.connect(DB_PATH)
    df = pd.read_sql_query("SELECT * FROM beneficios ORDER BY id DESC", con)
    con.close()
    if not df.empty:
        st.subheader("Bitácora de beneficios")
        st.dataframe(df.drop(columns=["id"]), use_container_width=True, hide_index=True)



def page_agua_productiva():
    st.title("Agua potable productiva · WASH + emprendimientos")
    st.markdown("""
    Este módulo orienta a los gerentes de campamentos transitorios para convertir la necesidad crítica de agua potable en una **capacidad comunitaria organizada**: brigadas WASH, microservicios de monitoreo, inventario de insumos, capacitación y emprendimientos emergentes que contribuyan al aseguramiento sostenible, sustentable y soportable del agua segura en el estado La Guaira.
    """)
    st.markdown("""
    <div class="notice"><b>Idea fuerza:</b> primero se asegura vida e higiene; luego se organiza talento; después se productiviza el conocimiento y se servitiza el agua segura como red comunitaria de valor público.</div>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3, tab4 = st.tabs(["Marco WASH", "Ruta de emprendimiento", "Lienzo de modelo", "Bitácora Agua Segura"])
    with tab1:
        st.subheader("WASH: agua, saneamiento e higiene")
        col1, col2 = st.columns([1,1])
        with col1:
            st.markdown("""
            **Pilares operativos**
            - **Agua:** acceso a agua segura para consumo, cocina y aseo.
            - **Saneamiento:** control de excretas, residuos, aguas servidas y puntos críticos.
            - **Higiene:** lavado de manos, inocuidad alimentaria y hábitos preventivos.

            **Ruta de transmisión fecal-oral:** heces/excretas → fluidos, campos, moscas, manos/dedos y alimentos → nuevo huésped susceptible.
            """)
        with col2:
            st.markdown("""
            **Barreras de interrupción**
            - Aislar excretas en la fuente.
            - Tratar agua por cloración/filtración.
            - Higiene de manos y alimentos.
            - Control de vectores y residuos.
            - Registro diario para alertas tempranas.
            """)
        st.subheader("Parámetros operativos sugeridos")
        parametros = pd.DataFrame([
            {"Variable":"Cloro residual libre (CRL)","Meta":"0.5 a 1.0 mg/L","Uso":"Agua potable en puntos de consumo; hasta 1.5 mg/L si ingresa cisterna recién tratada."},
            {"Variable":"Lecturas diarias","Meta":"Mínimo 3","Uso":"Tanque/cisterna, cocina/filtro y baños/estaciones."},
            {"Variable":"Sanitización de áreas comunes","Meta":"Solución clorada 0.1% / 1000 ppm","Uso":"2 a 3 veces al día en cocina, comedor, sanitarios, basura y aislamiento preventivo."},
            {"Variable":"Hortalizas/frutas","Meta":"50 ppm","Uso":"1 mL de cloro comercial por litro de agua; contacto 5 a 10 minutos y enjuague con agua potable."},
            {"Variable":"Utensilios/mesones","Meta":"200 ppm","Uso":"4 mL de cloro por litro de agua; contacto 3 a 5 minutos y secado al aire."},
            {"Variable":"Cocción segura","Meta":">70 °C","Uso":"Alimentos cocidos y servidos inmediatamente para evitar proliferación bacteriana."},
        ])
        st.dataframe(parametros, use_container_width=True, hide_index=True)

    with tab2:
        st.subheader("Ruta de transformación: necesidad → capacidad → emprendimiento")
        ruta = load_csv("ruta_emprendimiento_agua.csv")
        if not ruta.empty:
            st.dataframe(ruta, use_container_width=True, hide_index=True)
        st.markdown("""
        <div class="warn"><b>Sostenible, sustentable y soportable:</b> sostenible porque puede mantenerse en el tiempo; sustentable porque protege recursos y ambiente; soportable porque no sobrecarga a las personas, la infraestructura ni la logística del campamento.</div>
        """, unsafe_allow_html=True)
        st.subheader("Modelos de emprendimiento sugeridos")
        modelos = load_csv("modelos_agua.csv")
        if not modelos.empty:
            st.dataframe(modelos, use_container_width=True, hide_index=True)

    with tab3:
        st.subheader("Lienzo inteligentista de modelo de negocio / servicio")
        modelos = load_csv("modelos_agua.csv")
        default_model = modelos["Modelo"].tolist()[0] if not modelos.empty else "Brigada Agua Segura"
        with st.form("form_emprendimiento_agua"):
            c1, c2 = st.columns(2)
            campamento = c1.text_input("Campamento transitorio", "Campamento Demo La Guaira")
            nombre = c2.text_input("Nombre de la iniciativa", "Agua Segura La Guaira")
            modelo = st.selectbox("Modelo sugerido", modelos["Modelo"].tolist() if not modelos.empty else [default_model])
            producto_servicio = st.text_area("Producto / servicio mínimo viable", "Monitoreo diario de agua potable, cloración orientativa, punto de recarga segura, capacitación de familias y bitácora WASH.")
            talento = st.text_area("Talento del campamento a activar", "Químicos, enfermeros, docentes, líderes comunitarios, plomeros, técnicos, jóvenes voluntarios y manipuladores de alimentos.")
            aliado = st.text_input("Aliado clave", "Protección Civil / Salud / Hidrológica / universidades / donantes WASH")
            indicador = st.text_input("Indicador de éxito", "Lecturas CRL en rango, reducción de casos DDA, familias capacitadas y puntos de agua aptos")
            accion = st.text_area("Próxima acción en 72 horas", "Formar brigada, identificar puntos de agua, levantar inventario, preparar bitácora y realizar primera ronda de medición.")
            etapa = st.selectbox("Etapa", ["Idea", "Piloto", "Operando", "Escalable", "Red estadal"])
            ok = st.form_submit_button("Guardar iniciativa de agua potable")
        if ok:
            con = sqlite3.connect(DB_PATH)
            cur = con.cursor()
            cur.execute("""INSERT INTO emprendimientos_agua(fecha,campamento,nombre_iniciativa,modelo,producto_servicio,talento,aliado,indicador,proxima_accion,etapa)
            VALUES(?,?,?,?,?,?,?,?,?,?)""", (datetime.now().isoformat(timespec="seconds"),campamento,nombre,modelo,producto_servicio,talento,aliado,indicador,accion,etapa))
            con.commit(); con.close()
            st.success("Iniciativa guardada.")
        con = sqlite3.connect(DB_PATH)
        try:
            df = pd.read_sql_query("SELECT * FROM emprendimientos_agua ORDER BY id DESC", con)
        except Exception:
            df = pd.DataFrame()
        con.close()
        if not df.empty:
            st.subheader("Iniciativas registradas")
            st.dataframe(df.drop(columns=["id"]), use_container_width=True, hide_index=True)

    with tab4:
        st.subheader("Bitácora rápida de Agua Segura")
        st.markdown("Registre mediciones y observaciones para orientar acciones. Esta bitácora puede exportarse a CSV para impresión o reporte.")
        with st.form("wash_log"):
            c1,c2,c3,c4 = st.columns(4)
            fecha = c1.date_input("Fecha")
            punto = c2.selectbox("Punto", ["Tanque / cisterna", "Cocina / filtro", "Baños / estaciones", "Punto comunitario", "Otro"])
            crl = c3.number_input("CRL mg/L", min_value=0.0, max_value=5.0, value=0.5, step=0.1)
            turbidez = c4.selectbox("Turbidez", ["Clara", "Ligeramente turbia", "Turbia"])
            apto = st.selectbox("¿Apto para consumo?", ["Sí", "No", "Requiere verificación"])
            obs = st.text_area("Acción correctiva / observación", "")
            guardar = st.form_submit_button("Generar registro")
        if guardar:
            log = pd.DataFrame([{"Fecha":fecha, "Punto":punto, "CRL_mg_L":crl, "Turbidez":turbidez, "Apto":apto, "Observacion":obs}])
            st.dataframe(log, use_container_width=True, hide_index=True)
            st.download_button("Descargar registro CSV", log.to_csv(index=False).encode("utf-8-sig"), "bitacora_agua_segura.csv", "text/csv")
        st.subheader("Semáforo interpretativo")
        sem = pd.DataFrame([
            {"Condición":"CRL < 0.5 mg/L", "Lectura":"Bajo", "Acción":"No liberar consumo sin corrección/verificación; revisar cloración y fuente."},
            {"Condición":"CRL 0.5-1.0 mg/L", "Lectura":"Meta", "Acción":"Mantener monitoreo y registros."},
            {"Condición":"CRL > 1.5 mg/L", "Lectura":"Alto", "Acción":"Verificar dosificación y esperar/mezclar según criterio técnico competente."},
            {"Condición":"Turbia", "Lectura":"Riesgo", "Acción":"Filtrar/decantar/tratar y escalar a autoridad sanitaria."},
        ])
        st.dataframe(sem, use_container_width=True, hide_index=True)



def page_escuela_wash_flash():
    st.title("Escuela WASH Flash")
    st.markdown("""
    Curso de emergencia para formar rápidamente a dignificados, brigadas, monitores de agua y gerentes de centro en WASH: agua segura, saneamiento, higiene, bioseguridad, inocuidad y análisis rápido de agua. La intención no es dictar una charla decorativa: es activar personas capaces de ver el riesgo, ejecutar prácticas, registrar evidencia, enseñar a otros y abrir microservicios comunitarios de agua potable para La Guaira.
    """)
    st.markdown("""
    <div class="notice"><b>Didáctica de transferencia tecnocognitiva:</b> Ver -> Hacer -> Medir -> Registrar -> Corregir -> Enseñar -> Servir -> Emprender. Cada módulo debe terminar con una evidencia observable y una decisión operativa.</div>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
        "Ruta 3-1-3",
        "Aulas y prácticas",
        "Videos oficiales",
        "Casos reales",
        "Evaluación",
        "Microcredenciales",
        "Puente emprendedor"
    ])

    with tab1:
        st.subheader("Curso flash WASH para respuesta post-sismo")
        st.markdown("""
        La Escuela WASH Flash está diseñada para entrenar en tiempo récord sin perder rigor humanitario: primero protege la vida y corta rutas de infección; luego organiza brigadas; después convierte el aprendizaje en servicios comunitarios sostenibles.
        """)
        curso = load_csv("curso_wash_flash.csv")
        if not curso.empty:
            st.dataframe(curso, use_container_width=True, hide_index=True)
        st.subheader("Ruta 3-1-3")
        plan = load_csv("ruta_313_wash.csv")
        if not plan.empty:
            st.dataframe(plan, use_container_width=True, hide_index=True)
        st.markdown("""
        <div class="ok"><b>Resultado esperado:</b> una Brigada WASH mínima operando, bitácoras activas, puntos críticos identificados, personas preseleccionadas por competencias para análisis de agua y familias orientadas sobre agua segura, saneamiento e higiene.</div>
        """, unsafe_allow_html=True)
        st.subheader("Curso específico: análisis rápido de agua")
        analisis = load_csv("curso_analisis_agua.csv")
        if not analisis.empty:
            st.dataframe(analisis, use_container_width=True, hide_index=True)
        st.markdown("<div class='warn'><b>Puerta de seguridad:</b> hasta que exista personal certificado o con responsable técnico, el análisis local es orientación operativa; la liberación formal del agua requiere validación competente cuando haya sospecha microbiológica.</div>", unsafe_allow_html=True)

    with tab2:
        st.subheader("Aulas prácticas: aprender haciendo")
        practicas = load_csv("practicas_wash.csv")
        if not practicas.empty:
            st.dataframe(practicas, use_container_width=True, hide_index=True)
        st.subheader("Plan didáctico por estación")
        estaciones = load_csv("estaciones_wash_flash.csv")
        if not estaciones.empty:
            st.dataframe(estaciones, use_container_width=True, hide_index=True)
        st.markdown("""
        <div class="warn"><b>Regla de aprobación:</b> nadie aprueba por oír la inducción. Cada participante debe demostrar una práctica: lavado correcto de manos, registro de agua, sanitización, ruta fecal-oral o charla familiar.</div>
        """, unsafe_allow_html=True)

    with tab3:
        st.subheader("Tutoriales y fuentes verificables")
        st.caption("Los enlaces externos deben revisarse o descargarse previamente cuando haya conectividad. Priorice fuentes oficiales o humanitarias reconocidas.")
        vids = load_csv("tutoriales_wash.csv")
        if not vids.empty:
            for _, r in vids.iterrows():
                st.markdown(f"""
                <div class="card">
                  <h4>{r['Titulo']}</h4>
                  <p><b>Fuente:</b> {r['Fuente']} | <b>Tema:</b> {r['Tema']} | <b>Duración sugerida en clase:</b> {r.get('Momento_clase','Ver según conectividad')}</p>
                  <p><b>Uso didáctico:</b> {r['Uso_didactico']}</p>
                  <p><a href="{r['URL']}" target="_blank">Abrir video / recurso</a></p>
                </div>
                """, unsafe_allow_html=True)
        st.markdown("""
        <div class="notice"><b>Uso recomendado:</b> proyectar 1 video corto, ejecutar una práctica de 10-15 minutos y cerrar con una pregunta: ¿qué cambiaremos hoy en este campamento?</div>
        """, unsafe_allow_html=True)

    with tab4:
        st.subheader("Casos reales y discusión rápida")
        casos = load_csv("casos_wash.csv")
        if not casos.empty:
            st.dataframe(casos, use_container_width=True, hide_index=True)
        st.subheader("Guía de discusión 3S")
        st.markdown("""
        - <b>Situación:</b> ¿qué ocurrió y qué riesgo apareció?
        - <b>Significado:</b> ¿qué implica para agua, salud, convivencia y confianza?
        - <b>Solución:</b> ¿qué acción concreta puede hacer la brigada en 24 horas?
        """, unsafe_allow_html=True)

    with tab5:
        st.subheader("Registro de capacitación y medición de aprendizaje")
        with st.form("evaluacion_wash_flash_form"):
            c1, c2, c3 = st.columns(3)
            campamento = c1.text_input("Campamento transitorio", "Campamento Demo La Guaira")
            grupo = c2.text_input("Grupo / pabellón", "Brigada WASH - Pabellón A")
            modulo = c3.selectbox("Módulo", ["Inducción WASH", "Agua segura", "Saneamiento", "Higiene", "Bioseguridad e inocuidad", "Microservicio Agua Segura"])
            c1, c2, c3 = st.columns(3)
            participantes = c1.number_input("Personas formadas", min_value=0, value=12, step=1)
            entrada = c2.slider("Evaluación de entrada", 0, 100, 40)
            salida = c3.slider("Evaluación de salida", 0, 100, 80)
            practica = st.selectbox("Práctica demostrada", ["Lavado de manos", "Bitácora de agua", "Diagrama F", "Sanitización", "Inocuidad alimentaria", "Charla familiar", "Punto de agua segura"])
            microcredencial = st.selectbox("Microcredencial sugerida", ["Participante WASH", "Monitor de Agua Segura", "Promotor de Higiene", "Sanitizador de Áreas", "Manipulador Seguro de Alimentos", "Multiplicador WASH"])
            obs = st.text_area("Observaciones", "El grupo completó práctica y definió responsables por turno.")
            proxima = st.text_area("Próxima acción", "Repetir práctica en el pabellón, registrar evidencia y formar dos multiplicadores.")
            ok = st.form_submit_button("Guardar evaluación")
        if ok:
            con = sqlite3.connect(DB_PATH)
            cur = con.cursor()
            cur.execute("""INSERT INTO evaluaciones_wash_flash(fecha,campamento,grupo,modulo,participantes,evaluacion_entrada,evaluacion_salida,practica,microcredencial,observaciones,proxima_accion)
            VALUES(?,?,?,?,?,?,?,?,?,?,?)""", (datetime.now().isoformat(timespec="seconds"), campamento, grupo, modulo, int(participantes), float(entrada), float(salida), practica, microcredencial, obs, proxima))
            con.commit(); con.close()
            st.success("Evaluación WASH Flash guardada.")
        try:
            con = sqlite3.connect(DB_PATH)
            df = pd.read_sql_query("SELECT * FROM evaluaciones_wash_flash ORDER BY id DESC", con)
            con.close()
        except Exception:
            df = pd.DataFrame()
        if not df.empty:
            st.dataframe(df.drop(columns=["id"]), use_container_width=True, hide_index=True)
            if len(df) > 0:
                fig = px.bar(df.tail(10), x="grupo", y=["evaluacion_entrada", "evaluacion_salida"], barmode="group", title="Aprendizaje antes/después por grupo")
                st.plotly_chart(fig, use_container_width=True)

    with tab6:
        st.subheader("Microcredenciales operativas")
        creds = load_csv("microcredenciales_wash.csv")
        if not creds.empty:
            st.dataframe(creds, use_container_width=True, hide_index=True)
        st.markdown("""
        <div class="ok"><b>Idea fuerza:</b> las microcredenciales no son títulos académicos; son reconocimientos operativos para activar confianza, roles y responsabilidad comunitaria.</div>
        """, unsafe_allow_html=True)
        with st.form("compromiso_wash_form"):
            c1, c2, c3 = st.columns(3)
            campamento = c1.text_input("Campamento", "Campamento Demo La Guaira", key="comp_camp")
            brigada = c2.text_input("Brigada", "Brigada Agua Segura")
            responsable = c3.text_input("Responsable", "Líder de turno")
            compromiso = st.text_area("Compromiso", "Mantener bitácora de agua y reportar alertas cada día.")
            indicador = st.text_input("Indicador", "3 registros diarios completos y acciones correctivas documentadas")
            plazo = st.selectbox("Plazo", ["24 horas", "72 horas", "7 días", "30 días"])
            estado = st.selectbox("Estado", ["Pendiente", "En marcha", "Cumplido", "Bloqueado"])
            ok2 = st.form_submit_button("Guardar compromiso")
        if ok2:
            con = sqlite3.connect(DB_PATH)
            cur = con.cursor()
            cur.execute("""INSERT INTO compromisos_wash(fecha,campamento,brigada,compromiso,responsable,indicador,plazo,estado) VALUES(?,?,?,?,?,?,?,?)""",
                        (datetime.now().isoformat(timespec="seconds"), campamento, brigada, compromiso, responsable, indicador, plazo, estado))
            con.commit(); con.close()
            st.success("Compromiso guardado.")
        try:
            con = sqlite3.connect(DB_PATH)
            dfc = pd.read_sql_query("SELECT * FROM compromisos_wash ORDER BY id DESC", con)
            con.close()
        except Exception:
            dfc = pd.DataFrame()
        if not dfc.empty:
            st.dataframe(dfc.drop(columns=["id"]), use_container_width=True, hide_index=True)

    with tab7:
        st.subheader("De Escuela WASH a emprendimientos de agua potable")
        st.markdown("""
        La capacitación debe convertirse en capacidad instalada. El gerente no solo forma personas: crea una red de servicios comunitarios que protege salud, genera ocupación sana y prepara emprendimientos de agua potable para La Guaira.
        """)
        puente = load_csv("puente_emprendedor_wash.csv")
        if not puente.empty:
            st.dataframe(puente, use_container_width=True, hide_index=True)
        st.markdown("""
        <div class="notice"><b>Ruta de valor:</b> necesidad crítica -> brigada entrenada -> servicio comunitario -> paquete productivizado -> red intercampamentos -> ciudadanía inteligente.</div>
        """, unsafe_allow_html=True)


def page_calidad_agua():
    st.title("Calidad del agua y análisis rápido WASH")
    st.markdown("""
    Este módulo responde al ajuste del tutor: la calidad del agua no puede quedar escondida en el emprendimiento; debe nacer en el diagnóstico inicial, generar necesidad de Escuela WASH y orientar decisiones de liberación, restricción o corrección.
    """)
    last = last_diagnostic()
    if last:
        d = last["data"]
        c1,c2,c3,c4 = st.columns(4)
        crl = float(d.get("crl_consumo", 0))
        turb = float(d.get("turbidez_ntu", 0))
        lect = int(d.get("lecturas_dia", 0))
        cert = int(d.get("certificados_agua", 0))
        c1.metric("CRL diagnóstico", f"{crl:.1f} mg/L")
        c2.metric("Turbidez", f"{turb:.1f} NTU")
        c3.metric("Lecturas/día", lect)
        c4.metric("Certificados", cert)
        alerts=[]
        if crl < 0.2: alerts.append("CRL por debajo del mínimo operativo: no liberar consumo sin corrección/verificación.")
        elif crl < 0.5: alerts.append("CRL bajo: reforzar monitoreo y verificar cloración según criterio técnico.")
        if turb > 5: alerts.append("Turbidez por encima de 5 NTU: filtrar/decantar/tratar y escalar si persiste.")
        if int(d.get("coliformes_detectados",0)) == 1: alerts.append("Coliformes/sospecha/sin prueba confiable: activar restricción y autoridad sanitaria.")
        if lect < 3: alerts.append("Menos de 3 lecturas diarias: activar bitácora mínima 07:00, 11:00 y 16:00.")
        if cert <= 0: alerts.append("No hay personas certificadas: el curso rápido de análisis de agua es urgente.")
        if alerts:
            st.markdown("<div class='danger'><b>Alertas del último diagnóstico:</b><br>" + "<br>".join(["- "+a for a in alerts]) + "</div>", unsafe_allow_html=True)
        else:
            st.markdown("<div class='ok'><b>Sin alertas críticas de calidad de agua en el último diagnóstico.</b></div>", unsafe_allow_html=True)

    tab1, tab2, tab3, tab4 = st.tabs(["Bitácora", "Semáforo Esfera-WASH", "Habilitación", "Demanda de curso"])
    with tab1:
        st.subheader("Registro de análisis rápido de agua")
        with st.form("analisis_agua_v06"):
            c1,c2,c3 = st.columns(3)
            campamento = c1.text_input("Campamento", "Campamento Demo La Guaira")
            punto = c2.selectbox("Punto de muestreo", ["Tanque/cisterna", "Grifo cocina", "Filtro", "Baños/estaciones", "Punto comunitario", "Otro"])
            analista = c3.text_input("Analista / monitor", "")
            c1,c2,c3,c4 = st.columns(4)
            crl = c1.number_input("CRL mg/L", min_value=0.0, max_value=5.0, value=0.2, step=0.1)
            turb = c2.number_input("Turbidez NTU", min_value=0.0, max_value=100.0, value=5.0, step=0.5)
            coliformes = c3.selectbox("Coliformes", ["No detectado", "Detectado", "No evaluado"])
            certificacion = c4.selectbox("Estado del analista", ["No entrenado", "En entrenamiento", "Certificación pendiente", "Certificado"])
            apto_auto = "Sí" if 0.2 <= crl <= 1.5 and turb <= 5 and coliformes == "No detectado" and certificacion in ["Certificación pendiente", "Certificado"] else "No / requiere verificación"
            apto = st.selectbox("¿Apto para consumo?", [apto_auto, "Sí", "No", "Requiere verificación"], index=0)
            accion = st.text_area("Acción correctiva", "Registrar, comunicar al gerente del centro y repetir medición si está fuera de rango.")
            ok = st.form_submit_button("Guardar análisis")
        if ok:
            con = sqlite3.connect(DB_PATH); cur = con.cursor()
            cur.execute("""INSERT INTO analisis_agua(fecha,campamento,punto,crl,turbidez,coliformes,apto,analista,certificacion,accion_correctiva) VALUES(?,?,?,?,?,?,?,?,?,?)""", (datetime.now().isoformat(timespec="seconds"), campamento, punto, float(crl), float(turb), coliformes, apto, analista, certificacion, accion))
            con.commit(); con.close()
            st.success("Análisis de agua guardado.")
        con = sqlite3.connect(DB_PATH)
        df = pd.read_sql_query("SELECT * FROM analisis_agua ORDER BY id DESC", con)
        con.close()
        if not df.empty:
            st.dataframe(df.drop(columns=["id"]), use_container_width=True, hide_index=True)
            st.download_button("Descargar bitácora de agua", df.drop(columns=["id"]).to_csv(index=False).encode("utf-8-sig"), "bitacora_analisis_agua.csv", "text/csv")

    with tab2:
        st.subheader("Indicadores mínimos para decisión rápida")
        sem = pd.DataFrame([
            {"Indicador":"Dotación mínima", "Meta":"15 L/persona/día inicial", "Decisión":"Si baja: racionar con transparencia, solicitar suministro y priorizar consumo/cocina/higiene básica."},
            {"Indicador":"CRL punto consumo", "Meta":"0.2-0.5 mg/L como referencia de emergencia", "Decisión":"Si <0.2: no liberar sin corrección. Si >1.5: verificar dosificación."},
            {"Indicador":"Turbidez", "Meta":"<5 NTU", "Decisión":"Si >5: decantar/filtrar/tratar y escalar."},
            {"Indicador":"Coliformes fecales", "Meta":"Ausencia", "Decisión":"Si detecta/sospecha: restringir consumo y activar autoridad sanitaria."},
            {"Indicador":"Lecturas diarias", "Meta":"3 lecturas/día", "Decisión":"Tanque/cisterna, cocina/filtro y punto comunitario o sanitario."},
            {"Indicador":"Personal habilitado", "Meta":"Al menos 2 por turno", "Decisión":"Si no hay: activar Escuela WASH - curso análisis de agua."},
        ])
        st.dataframe(sem, use_container_width=True, hide_index=True)
        st.markdown("<div class='warn'><b>Precaución:</b> SITRA orienta decisiones; no sustituye confirmación de laboratorio ni autoridad sanitaria cuando hay sospecha microbiológica.</div>", unsafe_allow_html=True)

    with tab3:
        st.subheader("Habilitación para usar la evaluación de agua")
        habil = pd.DataFrame([
            {"Estado":"No entrenado", "Puede medir":"Solo acompaña", "Puede decidir":"No", "Uso de dato":"Orientativo y bajo supervisión."},
            {"Estado":"En entrenamiento", "Puede medir":"Sí, con tutor", "Puede decidir":"No", "Uso de dato":"Registro de práctica, no liberación formal."},
            {"Estado":"Certificación pendiente", "Puede medir":"Sí", "Puede decidir":"Solo con responsable técnico", "Uso de dato":"Puede apoyar plan de acción."},
            {"Estado":"Certificado", "Puede medir":"Sí", "Puede decidir":"Recomienda; gerente valida", "Uso de dato":"Impacta diagnóstico y plan."},
        ])
        st.dataframe(habil, use_container_width=True, hide_index=True)

    with tab4:
        st.subheader("Generar necesidad de curso de análisis de agua")
        with st.form("demanda_curso_form"):
            c1,c2,c3 = st.columns(3)
            campamento = c1.text_input("Campamento", "Campamento Demo La Guaira", key="dem_camp")
            personas_objetivo = c2.number_input("Personas objetivo", min_value=1, value=25, step=1)
            estado = c3.selectbox("Estado", ["Necesidad detectada", "Curso programado", "En curso", "Certificación pendiente", "Cerrado"])
            perfil = st.text_area("Perfil objetivo", "Personas con lectura básica, disciplina de registro, liderazgo comunitario, experiencia en cocina/salud/mantenimiento o interés técnico.")
            razon = st.text_area("Razón de necesidad", "El campamento requiere analizar agua de manera ágil porque la vida, salud, cocina, higiene y convivencia dependen del agua segura.")
            curso = st.selectbox("Curso requerido", ["Análisis rápido de agua", "Monitor WASH", "Promotor de higiene", "Gestión de bitácoras", "Microservicio Agua Segura"])
            proxima = st.text_area("Próxima acción", "Seleccionar participantes, ubicar facilitador técnico y abrir primer curso 3-1-3.")
            ok = st.form_submit_button("Guardar demanda")
        if ok:
            con = sqlite3.connect(DB_PATH); cur = con.cursor()
            cur.execute("""INSERT INTO demanda_curso_wash(fecha,campamento,personas_objetivo,perfil_objetivo,razon_necesidad,curso_requerido,estado,proxima_accion) VALUES(?,?,?,?,?,?,?,?)""", (datetime.now().isoformat(timespec="seconds"), campamento, int(personas_objetivo), perfil, razon, curso, estado, proxima))
            con.commit(); con.close()
            st.success("Demanda de curso guardada.")
        con = sqlite3.connect(DB_PATH)
        df = pd.read_sql_query("SELECT * FROM demanda_curso_wash ORDER BY id DESC", con)
        con.close()
        if not df.empty:
            st.dataframe(df.drop(columns=["id"]), use_container_width=True, hide_index=True)


def page_gerencia_centro():
    st.title("Entrenamiento del gerente del centro")
    st.markdown("""
    Este módulo prepara al gerente del centro y a su suplente para coordinar organizaciones, priorizar necesidades básicas y sostener la operación cuando el jefe principal deba ausentarse. El tutor fue claro: no basta con tener instituciones presentes; hay que gerenciar el centro.
    """)
    tab1, tab2, tab3 = st.tabs(["Ruta del gerente", "Mapa de organizaciones", "Plan 24-72 h"])
    with tab1:
        st.subheader("Curso mínimo para gerente y suplente")
        ruta = pd.DataFrame([
            {"Bloque":"Rol del gerente", "Duración":"30 min", "Producto":"Responsables por agua, salud, alimentación, convivencia, información y logística."},
            {"Bloque":"Diagnóstico SITRA", "Duración":"45 min", "Producto":"Primer diagnóstico con calidad del agua incluida."},
            {"Bloque":"Agua como prioridad vital", "Duración":"45 min", "Producto":"Decisión de Escuela WASH y análisis rápido de agua."},
            {"Bloque":"Coordinación interinstitucional", "Duración":"45 min", "Producto":"Matriz: quién provee, quién cocina, quién mide, quién reporta."},
            {"Bloque":"Suplencia y continuidad", "Duración":"30 min", "Producto":"Un jefe operativo y un suplente entrenado por área."},
            {"Bloque":"Productivización", "Duración":"45 min", "Producto":"Curso/servicio WASH replicable para otros centros."},
        ])
        st.dataframe(ruta, use_container_width=True, hide_index=True)
        st.markdown("<div class='notice'><b>Principio operativo:</b> cada necesidad básica debe tener responsable, suplente, indicador, evidencia y próxima acción.</div>", unsafe_allow_html=True)
    with tab2:
        st.subheader("Mapa rápido de organizaciones y roles")
        orgs = pd.DataFrame([
            {"Organización":"INN / Nutrición", "Rol posible":"menú, inocuidad, nutrición vulnerable", "Pregunta gerencial":"¿Qué aporta hoy y qué indicador reporta?"},
            {"Organización":"Mercal / alimentación", "Rol posible":"suministro de alimentos", "Pregunta gerencial":"¿Qué entrega, cuándo y a quién?"},
            {"Organización":"Salud / epidemiología", "Rol posible":"alertas DDA, piel, respiratorias", "Pregunta gerencial":"¿Quién valida alertas y protocolos?"},
            {"Organización":"Hidrológica / agua", "Rol posible":"suministro, cloración, cisternas", "Pregunta gerencial":"¿Quién garantiza fuente, volumen y calidad?"},
            {"Organización":"Protección Civil", "Rol posible":"seguridad, riesgo, evacuación", "Pregunta gerencial":"¿Cuál es el canal de escalamiento?"},
            {"Organización":"Universidad / facilitadores", "Rol posible":"curso análisis de agua y WASH", "Pregunta gerencial":"¿Quién certifica práctica y acompaña?"},
        ])
        st.dataframe(orgs, use_container_width=True, hide_index=True)
    with tab3:
        st.subheader("Plan gerencial 24-72 horas")
        with st.form("gerencia_centro_form"):
            c1,c2,c3 = st.columns(3)
            campamento = c1.text_input("Campamento", "Campamento Demo La Guaira", key="gc_camp")
            gerente = c2.text_input("Gerente del centro", "Walter / responsable designado")
            suplente = c3.text_input("Suplente", "")
            organizaciones = st.text_area("Organizaciones presentes", "INN, Mercal, Salud, Protección Civil, Hidrológica, comunidad, voluntarios")
            necesidad = st.selectbox("Necesidad prioritaria", ["Agua potable", "Salud", "Convivencia", "Alimentación", "Saneamiento", "Información", "Logística"])
            plan24 = st.text_area("Plan 24 h", "Diagnóstico SITRA, calidad del agua, responsables por área y bitácoras activas.")
            plan72 = st.text_area("Plan 72 h", "Curso análisis rápido de agua, brigadas por turno, suplencias y primera evaluación de indicadores.")
            estado = st.selectbox("Estado", ["Planificado", "En marcha", "Bloqueado", "Cumplido"])
            ok = st.form_submit_button("Guardar plan gerencial")
        if ok:
            con = sqlite3.connect(DB_PATH); cur = con.cursor()
            cur.execute("""INSERT INTO gerencia_centro(fecha,campamento,gerente,suplente,organizaciones,necesidad_prioritaria,plan_24h,plan_72h,estado) VALUES(?,?,?,?,?,?,?,?,?)""", (datetime.now().isoformat(timespec="seconds"), campamento, gerente, suplente, organizaciones, necesidad, plan24, plan72, estado))
            con.commit(); con.close()
            st.success("Plan gerencial guardado.")
        con = sqlite3.connect(DB_PATH)
        df = pd.read_sql_query("SELECT * FROM gerencia_centro ORDER BY id DESC", con)
        con.close()
        if not df.empty:
            st.dataframe(df.drop(columns=["id"]), use_container_width=True, hide_index=True)

def page_transiliencia_smart():
    st.title("Transiliencia y ciudadanía inteligente")
    last = last_diagnostic()
    if not last:
        st.info("Primero registre un diagnóstico rápido.")
        return
    trans = last["trans"]
    dims = pd.DataFrame({"Dimensión": list(trans["dimensiones"].keys()), "Puntaje": list(trans["dimensiones"].values())})
    c1,c2 = st.columns([1,2])
    c1.metric("Índice de transiliencia", trans["indice"], trans["etapa"])
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(r=dims["Puntaje"], theta=dims["Dimensión"], fill='toself', name='Transiliencia'))
    fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0,100])), showlegend=False, height=430)
    c2.plotly_chart(fig, use_container_width=True)
    st.markdown("""
    <div class="notice"><b>Ruta R-T:</b> R0 Emergencia → R1 Estabilización → R2 Resiliencia → T1 Transiliencia → T2 Ciudadanía inteligente.</div>
    """, unsafe_allow_html=True)
    st.dataframe(dims, use_container_width=True, hide_index=True)
    st.subheader("Del campamento a comunidad inteligente")
    st.markdown("""
    La lectura Smart City se operacionaliza sin tecnocentrismo: primero personas, organización, información confiable, redes, aprendizaje y gobernanza comunitaria. La tecnología es soporte; el ciudadano organizado es el verdadero sensor, actor y constructor de la nueva ciudad.
    """)


def page_i2e():
    st.title("I2E + A3D + 3S")
    st.markdown("Una pantalla para convertir presión, datos y contexto en decisiones ejecutables.")
    c1,c2,c3 = st.columns(3)
    with c1:
        st.subheader("I2E")
        intorno = st.text_area("Intorno: dentro del campamento", "Hacinamiento, agua limitada, conflictos leves, voluntarios activos.")
        entorno = st.text_area("Entorno: comunidad cercana", "Acceso parcialmente limitado, apoyo comunitario disponible.")
        extorno = st.text_area("Extorno: fuerzas externas", "Riesgo de lluvia, retraso en suministros, cooperación externa variable.")
    with c2:
        st.subheader("A3D")
        agil = st.text_area("Ágil: qué se hace hoy", "Priorizar agua, seguridad y comunicación.")
        digital = st.text_area("Digital: qué se registra", "Personas, recursos, riesgos, acciones, evidencias y beneficios.")
        disruptivo = st.text_area("Disruptivo/diferenciador", "Convertir el campamento en comunidad organizada con brigadas internas y mapa de talento.")
    with c3:
        st.subheader("3S")
        situacion = st.text_area("Situación", "Hay tensión en la distribución de alimentos.")
        significado = st.text_area("Significado", "Riesgo de conflicto, pérdida de confianza y desorden.")
        solucion = st.text_area("Solución", "Turnos visibles, comité comunitario y canal de quejas.")
    if st.button("Generar síntesis 3S"):
        st.markdown(f"""
        <div class="card"><h4>Síntesis situacional</h4>
        <p><b>Intorno:</b> {intorno}</p><p><b>Entorno:</b> {entorno}</p><p><b>Extorno:</b> {extorno}</p>
        <p><b>Decisión ágil:</b> {agil}</p><p><b>Evidencia digital:</b> {digital}</p><p><b>Cambio diferenciador:</b> {disruptivo}</p>
        <p><b>3S:</b> {situacion} → {significado} → {solucion}</p></div>
        """, unsafe_allow_html=True)


def page_dashboard():
    st.title("Dashboard")
    last = last_diagnostic()
    if not last:
        st.info("Primero registre un diagnóstico rápido.")
        return
    risks = last["risks"]
    trans = last["trans"]
    strategy = last["strategy"]
    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Riesgo promedio", round(risks["Puntaje"].mean(),1))
    c2.metric("Riesgo máximo", round(risks["Puntaje"].max(),1), risks.loc[risks["Puntaje"].idxmax(), "Categoría"])
    c3.metric("Transiliencia", trans["indice"], trans["etapa"])
    c4.metric("Madurez 12E", round(strategy["Madurez"].mean(),1) if not strategy.empty else "N/D")
    col1,col2 = st.columns(2)
    fig1 = px.pie(risks, names="Nivel", title="Distribución por nivel de riesgo")
    col1.plotly_chart(fig1, use_container_width=True)
    fig2 = px.bar(risks, x="Categoría", y="Puntaje", color="Nivel", title="Riesgo por categoría")
    fig2.update_layout(xaxis_tickangle=-35)
    col2.plotly_chart(fig2, use_container_width=True)
    if not strategy.empty:
        fig3 = px.bar(strategy, x="Código", y="Madurez", color="Nivel", title="Madurez de 12 elementos estratégicos")
        st.plotly_chart(fig3, use_container_width=True)
    hist = all_diagnostics()
    if len(hist) > 1:
        rows=[]
        for _, r in hist.iterrows():
            risks_i = pd.DataFrame(json.loads(r["riesgos_json"]))
            trans_i = json.loads(r["transiliencia_json"])
            strategy_i = pd.DataFrame(json.loads(r["estrategia_json"])) if r.get("estrategia_json") else pd.DataFrame()
            rows.append({"fecha": r["fecha"], "riesgo_promedio": risks_i["Puntaje"].mean(), "transiliencia": trans_i["indice"], "madurez_estrategica": strategy_i["Madurez"].mean() if not strategy_i.empty else None})
        trend = pd.DataFrame(rows).sort_values("fecha")
        fig = px.line(trend, x="fecha", y=["riesgo_promedio", "transiliencia", "madurez_estrategica"], title="Evolución histórica")
        st.plotly_chart(fig, use_container_width=True)


def create_docx_report(last):
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(10)
    doc.add_heading('Reporte SITRA-Campamentos', 0)
    doc.add_paragraph(f"Campamento transitorio: {last['campamento']}")
    doc.add_paragraph(f"Municipio: {last['municipio']}")
    doc.add_paragraph(f"Fecha: {last['fecha']}")
    doc.add_paragraph('Este reporte es un apoyo a la decisión y no sustituye a las autoridades competentes.')
    doc.add_heading('Mapa de riesgos', level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text='Categoría'; hdr[1].text='Puntaje'; hdr[2].text='Nivel'; hdr[3].text='Semáforo'
    for _, r in last['risks'].iterrows():
        cells = table.add_row().cells
        cells[0].text=str(r['Categoría']); cells[1].text=str(r['Puntaje']); cells[2].text=str(r['Nivel']); cells[3].text=str(r['Semáforo'])
    doc.add_heading('Madurez estratégica 12E', level=1)
    if not last['strategy'].empty:
        t2 = doc.add_table(rows=1, cols=3)
        h = t2.rows[0].cells
        h[0].text='Código'; h[1].text='Madurez'; h[2].text='Nivel'
        for _, r in last['strategy'].iterrows():
            c = t2.add_row().cells
            c[0].text=str(r['Código']); c[1].text=str(r['Madurez']); c[2].text=str(r['Nivel'])
    doc.add_heading('Plan de acción alineado', level=1)
    for _, r in last['plan'].iterrows():
        doc.add_paragraph(f"[{r['Horizonte']}] {r['Categoría']} | {r['Objetivo relacionado']} | {r['Elemento estratégico']} - {r['Acción recomendada']}")
    doc.add_heading('Transiliencia', level=1)
    doc.add_paragraph(f"Índice: {last['trans']['indice']} | Etapa: {last['trans']['etapa']}")
    doc.add_heading('Agua potable productiva', level=1)
    doc.add_paragraph('Este eje orienta brigadas WASH, productivizacion y servitizacion del agua segura para convertir la emergencia en capacidades comunitarias y emprendimientos sostenibles. La Escuela WASH Flash agrega curso breve, practicas, videos, registro de capacitacion y puente emprendedor.')
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def page_reportes():
    st.title("Reportes")
    last = last_diagnostic()
    if not last:
        st.info("Primero registre un diagnóstico rápido.")
        return
    st.write(f"Último diagnóstico: **{last['campamento']}** - {last['fecha']}")
    summary = {
        "campamento": last["campamento"],
        "fecha": last["fecha"],
        "riesgo_promedio": round(last["risks"]["Puntaje"].mean(),1),
        "riesgo_maximo": round(last["risks"]["Puntaje"].max(),1),
        "transiliencia": last["trans"]["indice"],
        "etapa": last["trans"]["etapa"],
        "madurez_12E": round(last["strategy"]["Madurez"].mean(),1) if not last["strategy"].empty else None
    }
    st.json(summary)
    st.download_button("Descargar riesgos CSV", last["risks"].to_csv(index=False).encode("utf-8-sig"), "riesgos_sitra_campamentos.csv", "text/csv")
    st.download_button("Descargar plan CSV", last["plan"].to_csv(index=False).encode("utf-8-sig"), "plan_sitra_campamentos.csv", "text/csv")
    if not last["strategy"].empty:
        st.download_button("Descargar 12E CSV", last["strategy"].to_csv(index=False).encode("utf-8-sig"), "madurez_12_elementos.csv", "text/csv")
    try:
        docx_bytes = create_docx_report(last)
        st.download_button("Descargar reporte Word", docx_bytes, "reporte_sitra_campamentos.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        st.warning(f"No se pudo generar Word. Detalle: {e}")


ensure_db()

st.sidebar.markdown(f"### {APP_NAME}")
st.sidebar.caption(APP_FULL)
page = st.sidebar.radio("Navegación", [
    "Inicio",
    "Base metodológica",
    "Diagnóstico rápido",
    "FODA Sistémica",
    "12 Elementos Estratégicos",
    "Objetivos Sistémicos",
    "Mapa de riesgos",
    "Plan de acción",
    "Beneficios + Océano Azul",
    "Agua Potable Productiva",
    "Calidad del agua",
    "Escuela WASH Flash",
    "Gerencia del centro",
    "Transiliencia + Smart City",
    "I2E + A3D + 3S",
    "Dashboard",
    "Reportes",
])
st.sidebar.markdown("---")
st.sidebar.caption(f"Prototipo local · {APP_VERSION} · Streamlit + SQLite + reglas expertas")

if page == "Inicio": page_inicio()
elif page == "Base metodológica": page_referencias()
elif page == "Diagnóstico rápido": page_diagnostico()
elif page == "FODA Sistémica": page_foda()
elif page == "12 Elementos Estratégicos": page_elementos()
elif page == "Objetivos Sistémicos": page_objetivos()
elif page == "Mapa de riesgos": page_mapa_riesgos()
elif page == "Plan de acción": page_plan_accion()
elif page == "Beneficios + Océano Azul": page_beneficios_oceano()
elif page == "Agua Potable Productiva": page_agua_productiva()
elif page == "Calidad del agua": page_calidad_agua()
elif page == "Escuela WASH Flash": page_escuela_wash_flash()
elif page == "Gerencia del centro": page_gerencia_centro()
elif page == "Transiliencia + Smart City": page_transiliencia_smart()
elif page == "I2E + A3D + 3S": page_i2e()
elif page == "Dashboard": page_dashboard()
elif page == "Reportes": page_reportes()

st.markdown('<div class="footer">SITRA-Campamentos v0.6.1 · Herramienta de apoyo a decisiones para campamentos transitorios. Uso responsable: complementar con autoridades competentes y protocolos oficiales.</div>', unsafe_allow_html=True)
